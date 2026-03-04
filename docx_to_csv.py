#!/usr/bin/env python3
"""Convert lowNiDiet summary tables .docx to CSV."""
import csv
import re
from pathlib import Path

from docx import Document

DOCX_PATH = Path.home() / "Downloads" / "lowNiDiet_r9.1.1_summaryTables.docx"
OUT_CSV = Path(__file__).parent / "lowNiDiet_r9.1.1_summaryTables.csv"

HEADER = [
    "Category",
    "Food or category (serving)",
    "Number of sources",
    "Mean Ni (ug/serving)",
    "Stddev (ug/serving)",
    "Min Ni (ug/serving)",
    "Max Ni (ug/serving)",
]


def cell_text(cell):
    return " ".join(p.text.strip() for p in cell.paragraphs).strip().replace("\n", " ")


def is_header_row(cells):
    if len(cells) < 2:
        return False
    c1 = (cells[1] or "").strip().lower()
    if c1 in ("number of", "sources") or "number of sources" in " ".join(cells).lower():
        return True
    if cells[0] == "Food or category (serving)" or cells[0] == "Dairy products and substitutes":
        if "sources" in c1 or "mean" in (cells[2] or "").lower():
            return True
    return False


def looks_numeric(s):
    if not s or not s.strip():
        return False
    return bool(re.match(r"^[\d.]+$", s.strip()))


def main():
    doc = Document(DOCX_PATH)
    rows_out = [HEADER]
    current_category = ""

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [cell_text(c) for c in row.cells]
            if not any(cells):
                continue

            # Detect category from first row of table (section header)
            if i == 0 and cells[0] and not looks_numeric(cells[1] or "x"):
                if "sources" in (cells[1] or "").lower() or (len(cells) > 2 and "mean" in (cells[2] or "").lower()):
                    current_category = cells[0]
                    continue  # skip header row
                # else: first row is data, keep current_category

            # Skip repeated header rows
            if is_header_row(cells):
                if cells[0] and not looks_numeric(cells[0]):
                    current_category = cells[0]
                continue

            # Pad to 6 cells, then prepend category
            while len(cells) < 6:
                cells.append("")
            row_data = [current_category] + cells[:6]
            rows_out.append(row_data)

    with open(OUT_CSV, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in rows_out:
            writer.writerow(row[:7])

    print(f"Wrote {len(rows_out) - 1} data rows (+ 1 header) to {OUT_CSV}")


if __name__ == "__main__":
    main()
